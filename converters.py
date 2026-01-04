import os
from moviepy import VideoFileClip, AudioFileClip
from PIL import Image
from docx import Document

conversion_options = {
    "mp4": ["avi", "mov", "mkv", "wmv", "flv", "webm", "ogv", "gif", "mpeg", "m4v", "3gp", "vob"],
    "avi": ["mp4", "mov", "mkv", "wmv", "flv", "webm", "ogv", "gif", "mpeg", "m4v", "3gp", "vob"],
    "mov": ["mp4", "avi", "mkv", "wmv", "flv", "webm", "ogv", "gif", "mpeg", "m4v", "3gp", "vob"],
    "mkv": ["mp4", "avi", "mov", "wmv", "flv", "webm", "ogv", "gif", "mpeg", "m4v", "3gp", "vob"],
    "wmv": ["mp4", "avi", "mov", "mkv", "flv", "webm", "ogv", "gif", "mpeg", "m4v", "3gp", "vob"],
    "flv": ["mp4", "avi", "mov", "mkv", "wmv", "webm", "ogv", "gif", "mpeg", "m4v", "3gp", "vob"],
    "webm": ["mp4", "avi", "mov", "mkv", "wmv", "flv", "ogv", "gif", "mpeg", "m4v", "3gp", "vob"],
    "ogv": ["mp4", "avi", "mov", "mkv", "wmv", "flv", "webm", "gif", "mpeg", "m4v", "3gp", "vob"],
    "gif": ["mp4", "avi", "mov", "mkv", "wmv", "flv", "webm", "ogv", "mpeg", "m4v", "3gp", "vob"],
    "mpeg": ["mp4", "avi", "mov", "mkv", "wmv", "flv", "webm", "ogv", "gif", "m4v", "3gp", "vob"],
    "m4v": ["mp4", "avi", "mov", "mkv", "wmv", "flv", "webm", "ogv", "gif", "mpeg", "3gp", "vob"],
    "3gp": ["mp4", "avi", "mov", "mkv", "wmv", "flv", "webm", "ogv", "gif", "mpeg", "m4v", "vob"],
    "vob": ["mp4", "avi", "mov", "mkv", "wmv", "flv", "webm", "ogv", "gif", "mpeg", "m4v", "3gp"],
    "mp3": ["wav", "ogg", "aac", "flac", "m4a"],
    "wav": ["mp3", "ogg", "aac", "flac", "m4a"],
    "ogg": ["mp3", "wav", "aac", "flac", "m4a"],
    "aac": ["mp3", "wav", "ogg", "flac", "m4a"],
    "flac": ["mp3", "wav", "ogg", "aac", "m4a"],
    "m4a": ["mp3", "wav", "ogg", "aac", "flac"],
    "jpg": ["png", "bmp", "gif", "tiff", "webp"],
    "png": ["jpg", "bmp", "gif", "tiff", "webp"],
    "bmp": ["jpg", "png", "gif", "tiff", "webp"],
    "gif": ["jpg", "png", "bmp", "tiff", "webp"],
    "tiff": ["jpg", "png", "bmp", "gif", "webp"],
    "webp": ["jpg", "png", "bmp", "gif", "tiff"],
    "pdf": ["docx", "txt"],
    "docx": ["pdf", "txt"],
    "txt": ["pdf", "docx"],
    "zip": ["tar"],
    "tar": ["zip"]
}

video_formats = ["mp4", "avi", "mov", "mkv", "wmv", "flv", "webm", "ogv", "gif", "mpeg", "m4v", "3gp", "vob"]
audio_formats = ["mp3", "wav", "ogg", "aac", "flac", "m4a"]
image_formats = ["jpg", "png", "bmp", "gif", "tiff", "webp"]
document_formats = ["pdf", "docx", "txt"]
archive_formats = ["zip", "tar"]

def convert_video(input_path, output_directory, output_format):
    input_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = f"{output_directory}/{input_name}.{output_format}"

    video_clip = VideoFileClip(input_path)
    video_clip.write_videofile(output_path, codec="libx264", audio_codec="aac")

def convert_audio(input_path, output_directory, output_format):
    input_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = f"{output_directory}/{input_name}.{output_format}"

    audio_clip = AudioFileClip(input_path)
    audio_clip.write_audiofile(output_path)

def convert_image(input_path, output_directory, output_format):
    input_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = f"{output_directory}/{input_name}.{output_format}"

    with Image.open(input_path) as img:
        img.save(output_path, format=output_format.upper())

def convert_document(input_path, output_directory, output_format):
    # Only simple docx <-> txt support for now
    input_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = f"{output_directory}/{input_name}.{output_format}"

    if input_path.endswith(".docx") and output_format == "txt":
        doc = Document(input_path)
        with open(output_path, "w", encoding="utf-8") as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")
    elif input_path.endswith(".txt") and output_format == "docx":
        doc = Document()
        with open(input_path, "r", encoding="utf-8") as f:
            for line in f:
                doc.add_paragraph(line.strip())
        doc.save(output_path)
    else:
        raise ValueError("Unsupported document conversion")

def convert_archive(input_path, output_directory, output_format):
    # Only zip <-> tar support (placeholder)
    input_name = os.path.splitext(os.path.basename(input_path))[0]
    output_path = f"{output_directory}/{input_name}.{output_format}"

    if input_path.endswith(".zip") and output_format == "tar":
        import zipfile, tarfile
        with zipfile.ZipFile(input_path, 'r') as zip_ref:
            extract_folder = f"{output_directory}/{input_name}_extracted"
            zip_ref.extractall(extract_folder)
            with tarfile.open(output_path, "w") as tar_ref:
                tar_ref.add(extract_folder, arcname=os.path.basename(extract_folder))
    elif input_path.endswith(".tar") and output_format == "zip":
        import tarfile, zipfile
        with tarfile.open(input_path, 'r') as tar_ref:
            extract_folder = f"{output_directory}/{input_name}_extracted"
            tar_ref.extractall(extract_folder)
            with zipfile.ZipFile(output_path, "w") as zip_ref:
                for root, _, files in os.walk(extract_folder):
                    for file in files:
                        file_path = os.path.join(root, file)
                        zip_ref.write(file_path, arcname=os.path.relpath(file_path, extract_folder))
    else:
        raise ValueError("Unsupported archive conversion")
